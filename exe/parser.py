from string import printable
import docx


class parser:
    file = None
    main_information = {'code': '', 'name': '', 'direction': '', 'profile': ''}
    competences = []
    developers = []
    themes = []

    def __init__(self, filename):
        self.file = docx.Document(filename)
        self.go_parse()

    def go_parse(self):
        self.find_competences()
        self.find_developers()
        self.find_main_information()
        self.find_themes()

    def find_main_information(self):
        self.main_information["code"], self.main_information["name"] = self.file.tables[1].rows[0].cells[3].text.split(
            " ", 1)
        self.main_information["direction"] = self.file.tables[1].rows[2].cells[1].text
        self.main_information["profile"] = self.file.tables[1].rows[4].cells[4].text

    def print_inform(self):
        for info in self.main_information:
            print(self.main_information[info])
        print()
        print()
        for competence in self.competences:
            for key in competence:
                if key == "indicators":
                    for item in competence[key]:
                        print(item[0])
                        print(item[1])
                print(competence[key])
            print()

    def get_main_information(self):
        return self.main_information

    def get_competences(self):
        return self.competences

    def get_developers(self):
        return self.developers

    def print(self):
        for row in self.file.tables[3].rows:
            for cell in row.cells:
                print(cell.text)
            print()

    def find_competences(self):
        buf = {'index': '', 'content': '', 'indicators': []}
        for row in self.file.tables[3].rows[1:]:
            index, content = row.cells[0].text.split(" ", 1)
            if buf['index'] == '' and buf['content'] == '':
                buf['index'] = index
                buf['content'] = content
            elif buf['index'] != index:
                self.competences.append(buf)
                buf = {'index': '', 'content': '', 'indicators': []}
            buf_indicator = (row.cells[1].text, row.cells[2].text)
            buf['indicators'].append(buf_indicator)
        self.competences.append(buf)

    def find_developers(self):
        buf = {'post': '', 'name': ''}
        for i in range(len(self.file.tables) - 1, 0, -1):
            table = self.file.tables[i]
            if len(table.rows[0].cells) == 5:
                row = table.rows[0]
                buf['post'] = row.cells[2].text
                buf['name'] = row.cells[4].text
                self.developers.append(buf)
                buf = {'post': '', 'name': ''}
            else:
                break

    def find_themes(self):
        buf = {'index': '', 'name': '', }
        for f_table in self.file.tables:
            if "Раздел дисциплины/темы" in f_table.rows[0].cells[1].text:
                table = f_table
                break
        for row in table.rows[3:-1]:
            if row.cells[0].text:
                if row.cells[0].text[-1] == '.' or not ('.' in row.cells[0].text):
                    buf['index'], buf['name'] = row.cells[0].text, row.cells[1].text
                    if buf['index'][-1] == '.':
                        buf['index'] = buf['index'][:-1]
                    if buf['index'] and buf['name']:
                        self.themes.append(buf)
                    buf = {'index': '', 'name': ''}

    def get_themes(self):
        return self.themes

    def get_idks(self):
        idks = []
        for compitence in self.competences:
            for item in compitence["indicators"]:
                idks.append((item[0], item[1]))
        return idks

    def clear(self):
        self.main_information.clear()
        self.competences.clear()
        self.developers.clear()
        self.themes.clear()
