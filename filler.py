from msilib.schema import tables
from textwrap import indent
from tkinter.messagebox import NO
from docx.enum.text import WD_COLOR_INDEX as colors
from turtle import position, width
import copy

from pip import main
from parser import parser
import docx


class filler:
	table_width = 5939790
	doc = None

	def __init__(self, filename):
		self.doc = docx.Document(filename)
	
	def fill_main_information(self, main_information):
		buf = ""
		first_run = None
		blue_run = None
		for paragraph in self.doc.paragraphs:
			for run in paragraph.runs:
				if run.font.highlight_color == colors.BLUE:
					if not(blue_run):
						blue_run = run
						run.font.highlight_color = colors.AUTO
					run.text = ""
				if run.font.highlight_color == colors.BRIGHT_GREEN:
					buf += run.text
					if (not(first_run)):
						first_run = run
					run.text = ""
				if run.font.highlight_color != colors.BRIGHT_GREEN and buf != "":
					first_run.text = main_information[buf[4:]]
					first_run.font.highlight_color = colors.AUTO
					buf = ""
					first_run = None

		for form in main_information['forms'][:-1]:
			blue_run.text += form.lower() + ", "
		blue_run.text += main_information['forms'][-1].lower()

		for table in self.doc.tables:
			for row in table.rows:
				for cell in row.cells:
					for paragraph in cell.paragraphs:
						for run in paragraph.runs:
							if run.font.highlight_color == colors.BRIGHT_GREEN:
								buf += run.text
								if (not(first_run)):
									first_run = run
								run.text = ""
							if run.font.highlight_color != colors.BRIGHT_GREEN and buf != "":
								first_run.text = main_information[buf[4:]]
								first_run.font.highlight_color = colors.AUTO
								buf = ""
								first_run = None

	def fill_competitions(self, competitions):
		for table in self.doc.tables:
			for row in table.rows:
				for cell in row.cells:
					if "Индекс" in cell.text:
						current_table = table
						break

		table = current_table
		for j in range(len(competitions)):
			self.make_new_competition_section(table, len(competitions[j]['indicators']))
			n = 0
			for i in range(len(table.rows) - len(competitions[j]['indicators']), len(table.rows)):
				table.rows[i].cells[0].text = str(j + 1)
				table.rows[i].cells[1].text = competitions[j]["index"]
				table.rows[i].cells[2].text = competitions[j]["content"]
				table.rows[i].cells[3].text = competitions[j]["indicators"][n][0]
				table.rows[i].cells[4].text = competitions[j]["indicators"][n][1]
				n += 1

	def make_new_competition_section(self, table, size):
		for i in range(size):
			table.add_row()
			if i > 0:
				table.rows[-1].cells[0].merge(table.rows[-2].cells[0])
				table.rows[-1].cells[1].merge(table.rows[-2].cells[1])
				table.rows[-1].cells[2].merge(table.rows[-2].cells[2])

	def fill_developers(self, developers):
		table = self.doc.tables[3]

		for i in range(len(self.doc.paragraphs)):
			paragraph = self.doc.paragraphs[i]
			if "developers" in paragraph.text:
				paragraph.text = ""
				id = i
		table.rows[0].cells[1].text = developers[0]["name"] + ", " + developers[0]["post"]

		if len(developers) > 1:
			position = 4
			for i in range(1, len(developers)):
				self.copy_table(table, self.doc.paragraphs[id])
				self.doc.tables[position].rows[0].cells[1].text = developers[i]["name"] + ", " + developers[i]["post"]
				position += 1
				id += 3

	def copy_table(self, table, paragraph):
		paragraph.text = ""
		paragraph._p.addnext(copy.deepcopy(table._tbl))

	def fill_themes(self, themes):
		
		for f_table in self.doc.tables:
			if "Тема или раздел" in f_table.rows[0].cells[0].text:
				table = f_table

		for item in themes:
			table.add_row()
			table.rows[-1].cells[0].text = item['index'] + " " + item['name']

	def save(self, filename):
		self.doc.save(filename)
		
a = filler("docx/shablon_test.docx")
b = parser("docx/rpd2.docx")
b.find_themes()
a.fill_developers(b.get_developers())
a.fill_competitions(b.get_comprtitions())
a.fill_themes(b.get_themes())
a.save("docx/new.docx")
