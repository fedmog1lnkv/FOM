from turtle import width
import docx


class table_of_competitions:

	def __init__(self, filename):
		doc = docx.Document(filename)
		tables = doc.tables[0]
	
	def add_competition(self, competition):
		print("хуй")

table_width = 5939790

doc = docx.Document("shablone.docx")

tables = [table for table in doc.tables]

print(tables[0])


print(tables[0].rows[1].cells[4].text)
print(tables[0].rows[2].cells[4].text)
print(tables[0].rows[3].cells[4].text)

doc.save("new_table.docx")
