from string import printable
import docx


class parser:
	file = None
	main_information = {'code': '', 'name' : '', 'direction' : '', 'profile' : ''}
	competitions = []
	developers = []
	themes = []

	def __init__(self, filename):
		self.file = docx.Document(filename)
		self.go_parse()

	def go_parse(self):
		self.main_information["code"], self.main_information["name"] = self.file.tables[1].rows[0].cells[3].text.split(" ", 1)
		self.main_information["direction"] = self.file.tables[1].rows[2].cells[1].text
		self.main_information["profile"] = self.file.tables[1].rows[4].cells[4].text
		self.find_competitions()
		self.find_developers()

		

	def print_inform(self):
		for info in self.main_information:
			print(self.main_information[info])
		print()
		print()
		for competition in self.competitions:
			for key in competition:
				if key == "indicators":
					for item in competition[key]:
						print(item[0])
						print(item[1])
				print(competition[key])
			print()

	def get_main_information(self):
		return self.main_information
		
	def get_comprtitions(self):
		return self.competitions
	
	def get_developers(self):
		return self.developers
		
	def print(self):
		for row in self.file.tables[3].rows:
			for cell in row.cells:
				print(cell.text)
			print()

	def find_competitions(self):
		buf = {'index' : '', 'content' : '', 'indicators' : []}
		for row in self.file.tables[3].rows[1:]:
			index, content = row.cells[0].text.split(" ", 1)
			if buf['index'] == '' and buf['content'] == '':
				buf['index'] = index
				buf['content'] = content
			elif buf['index'] != index:
				self.competitions.append(buf)
				buf = {'index' : '', 'content' : '', 'indicators' : []}
			buf_indicator = (row.cells[1].text, row.cells[2].text)
			buf['indicators'].append(buf_indicator)
		self.competitions.append(buf)

	def find_developers(self):
		buf = {'post' : '', 'name' : ''}
		for i in range(len(self.file.tables) - 1, 0, -1):
			table = self.file.tables[i]
			if len(table.rows[0].cells) == 5:
				row = table.rows[0]
				buf['post'] = row.cells[2].text
				buf['name'] = row.cells[4].text
				self.developers.append(buf)
				buf = {'post' : '', 'name' : ''}
			else:
				break

	def find_themes(self):
		buf = {'index' : '', 'name' : ''}
		for f_table in self.file.tables:
			if "Раздел дисциплины/темы" in f_table.rows[0].cells[1].text:
				table = f_table
				break
		for row in table.rows[3:-1]:
			buf['index'], buf['name'] = row.cells[0].text, row.cells[1].text
			if buf['index'] and buf['name']:
				self.themes.append(buf)
			buf = {'index' : '', 'name' : ''}
		
	def get_themes(self):
		return self.themes
		
		
